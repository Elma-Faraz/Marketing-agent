import os
from openai import AzureOpenAI
import json
import requests  # Import requests to download the image
from docx import Document  # To read Word files
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING,WD_ALIGN_PARAGRAPH 
from docx.shared import Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import docx
import shutil  # Import shutil for deleting folders
from PIL import Image
from io import BytesIO
from copy import deepcopy
from crewai import Crew, Process
from agents import create_post_writer_agent, create_validator_agent, create_prompt_agent, edited_content_validator_agent, content_regenerator_agent
from tasks import create_blog_write_task, create_validator_task, create_image_generator_task, create_prompt_task, edited_content_validator_task, regenerate_content_task
from constant import LANGUAGE_NAME_TO_CODE, default_styles
from dotenv import load_dotenv
import re

load_dotenv()
OPENAI_API_KEY = os.getenv("AZURE_API_KEY")
OPEN_API_ENDPOINT = os.getenv("AZURE_ENDPOINT")
AZURE_TRANSLATOR_KEY = os.getenv("AZURE_TRANSLATOR_KEY")
AZURE_TRANSLATOR_ENDPOINT = os.getenv("AZURE_TRANSLATOR_ENDPOINT")
AZURE_REGION = os.getenv("AZURE_REGION")
API_VERSION = '3.0'

#Method to process file and extract text
def process_uploaded_file(file):
    try:
        if file is not None:
            # Save the uploaded file to a temporary location
            temp_file_path = os.path.join("temp", file.name)
            with open(temp_file_path, "wb") as f:
                f.write(file.getbuffer())
            return temp_file_path
    except Exception as e:
        print(f"Error processing uploaded file: {e}")
    return ""

#method to call Crew() class
def crew_result(agents:list,tasks:list):
    try:
        if not agents or not tasks:
            raise ValueError("Agents and tasks cannot be empty.")
        crew = Crew(
                agents=agents,
                tasks=tasks,
            )
        result = crew.kickoff()
        
        # Extract the content from the result object
        content = None
        if result.json_dict and 'content' in result.json_dict:
            content = result.json_dict['content']
        elif result.raw:
            content = result.raw
        elif result.tasks_output:
            for task_output in result.tasks_output:
                if hasattr(task_output, 'output') and 'content' in task_output.output:
                    content = task_output.output['content']
                    break
        if not content:
            content = "No content generated."
        return content
    except Exception as e:
        print(f"Error in crew_result: {e}")
        return "No content generated."

#Method to create image
def create_image(prompt,image_name):
    try:  
        # If prompt is None or empty, return None
        if not prompt: 
            print("Prompt is empty or None.")
            return None
        
        # Ensure the images directory exists
        if not os.path.exists("images"):
            os.makedirs("images")
        
        client = AzureOpenAI(
        api_version="2024-02-01",
        azure_endpoint=OPEN_API_ENDPOINT,
        api_key=OPENAI_API_KEY,
        )

        result = client.images.generate(
            model="dall-e-3",  # the name of your DALL-E 3 deployment
            prompt=prompt,
            n=1,
            size="1024x1024",
        )

        # Extract the image URL
        image_url = json.loads(result.model_dump_json())['data'][0]['url']
        
        image_name = image_name

        # Download the image and save it to the 'images' folder
        response = requests.get(image_url)
        if response.status_code == 200:
            image_path = os.path.join("images", image_name)
            with open(image_path, "wb") as file:
                file.write(response.content)
            return image_path
        else:
            print(f"Failed to download the image. Status code: {response.status_code}")
            return image_path
    except Exception as e:
        print(f"Error generating image: {e}")
        return None

#method to update additional instructions with confirgurations
def additional_instructions(configurations: dict):
    try:
        if configurations is None:
            return ""
        additional_instructions = f"""
        The content type should be : {configurations["content_type"]}
        Write the content with length: {configurations["length"]}
        The tone of the content should be: {configurations["tone"]}
        The target audience of the content should be: {configurations["target_audience"]}
        The content should be for technicality level: {configurations["technicality"]}
        """
        return additional_instructions
    except Exception as e:
        print(f"Error during additional_instructions")
        return ""

# method to combine all the instructions into one
def combined_inputs(guidelines, instructions, additional_instructions):
    try:
        print("*****guidelines",guidelines)
        print("*****instructions",instructions)
        print("*****additional_instructions",additional_instructions)
        combined_input = f"""
        1. **Guidelines** : {guidelines}
        2. **Instructions** : {instructions}
        3. **Additional Instructions** : {additional_instructions}
        """
        return combined_input
    except Exception as e:
        print(f"Error during combined_inputs")
        return f"""
        1. **Guidelines** : None
        2. **Instructions** : None
        3. **Additional Instructions** : None
        """

#Method to generate instruction report
def generate_instruction_report(guidelines,configurations:dict):
    try:
        # create the prompt using configuration given by user
        config_instructions = additional_instructions(configurations)
        
        #Executing Instructions Conflicts task
        validator_agent = create_validator_agent(guidelines, config_instructions)
        validator_task = create_validator_task(guidelines, config_instructions)
        instruction_conflicts_report = crew_result([validator_agent],[validator_task])
        
        return instruction_conflicts_report
    except Exception as e:
        print(f"Error during generate_instruction_report: {e}")
        return "No report generated."

#Method to generate final output i.e., generated content, image_path, downloaded file path
def final_output(guidelines, instructions, configurations:dict, filename, output_language, generate_image_flag, image_name):
    try:
        # get output language code
        output_language_code = get_lang_code(output_language)
        
        # detect guidelines and instructions language
        detect_guideline_lang = detect_language(guidelines)
        detect_instructions_lang = detect_language(instructions)
        
        #translate guidelines and instructions to english if required
        if detect_guideline_lang != "en":
            translated_guidelines = translate_text(guidelines, to_lang="en", from_lang=detect_guideline_lang)
            guidelines = translated_guidelines
            print("*******translated guidelines: ",guidelines)
        
        if detect_instructions_lang != "en":
            translated_instructions = translate_text(instructions, to_lang="en", from_lang=detect_instructions_lang)
            instructions = translated_instructions
            print("*******translated instructions: ",instructions)
        
        # create the prompt using configuration given by user
        config_instructions = additional_instructions(configurations)
        
        #Combining all the inputs to pass to agent
        combined_input = combined_inputs(guidelines, instructions, config_instructions)
        
        # Executing blog writing task
        writer_agent = create_post_writer_agent(combined_input)
        writer_task = create_blog_write_task(combined_input)
        
        agents = [writer_agent]
        tasks = [writer_task]
        
        generated_post = crew_result(agents, tasks)
        generated_content_for_image = generated_post
        
        if output_language != "en":
            translated_post = translate_text(generated_post, to_lang=output_language_code, from_lang="en")
            generated_post = translated_post
        
        dir_name = "output_files"
        if not os.path.exists(dir_name):
            os.makedirs(dir_name)
        
        output_path = os.path.join(dir_name, filename)
        #style_guide_path = "input_guide/styling_guide.docx"
        
        format_document(guidelines, generated_post, output_path)
        print(
            f"Document saved at------------------------------------------------: {output_path}"
        )
        
        #image generation
        image = image_integration(output_path, generate_image_flag, image_name, generated_content_for_image, guidelines, instructions)

        final_output = {
            "Generated Post": generated_post,
            "Output File": output_path,
            "Imge": image
        }
        
        return final_output
    except Exception as e:
        print(f"Error during final_output: {e}")
        return {
            "Generated Post": "No post generated",
            "Output File": None,
            "Image":None
        }

# Function to read the styling guide from a Word file
def read_styling_guide(styling_guide):
    styling_guide_lines = styling_guide.splitlines()

    styles = {}

    try:
        # Define regex patterns for extracting formatting details
        font_pattern = re.compile(r"Font:\s*([\w\s]+)")
        size_pattern = re.compile(r"Size:\s*(\d+)")
        bold_pattern = re.compile(r"Bold:\s*(Yes|No)")
        italic_pattern = re.compile(r"Italic:\s*(Yes|No)")
        alignment_pattern = re.compile(r"Alignment:\s*(\w+)")
        color_pattern = re.compile(r"Color:\s*([\w\s]+|RGB\(\d+,\s*\d+,\s*\d+\))")
        spacing_pattern = re.compile(r"Spacing After:\s*(\d+)")
        underline_pattern = re.compile(r"Underline:\s*(Yes|No)")

        # Iterate through each paragraph in the styling guide
        current_style = None
        for line in styling_guide_lines:
            text = line.strip()
            if not text:
                continue

            # Check if the paragraph is a style name
            if text in ["Title", "Heading 1", "Heading 2", "Paragraph", "Bullet", "Hyperlink"]:
                current_style = text
                styles[current_style] = {}
                continue

            # Extract formatting details for the current style
            if current_style:
                font_match = font_pattern.search(text)
                size_match = size_pattern.search(text)
                bold_match = bold_pattern.search(text)
                italic_match = italic_pattern.search(text)
                alignment_match = alignment_pattern.search(text)
                color_match = color_pattern.search(text)
                spacing_match = spacing_pattern.search(text)
                underline_match = underline_pattern.search(text)

                # Parse the extracted details and add them to the dictionary
                if font_match:
                    styles[current_style]["font_name"] = font_match.group(1).strip()
                if size_match:
                    styles[current_style]["size"] = int(size_match.group(1))
                if bold_match:
                    styles[current_style]["bold"] = bold_match.group(1).strip().lower() == "yes"
                if italic_match:
                    styles[current_style]["italic"] = italic_match.group(1).strip().lower() == "yes"
                if alignment_match:
                    alignment = alignment_match.group(1).strip().lower()
                    alignment_map = {
                        "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                        "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                        "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                        "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                    }
                    styles[current_style]["alignment"] = alignment_map.get(alignment, WD_PARAGRAPH_ALIGNMENT.LEFT)
                if color_match:
                    color_value = color_match.group(1).strip()
                    if color_value.lower().startswith("rgb"):
                        # Parse RGB values from the format RGB(r, g, b)
                        rgb_match = re.match(r"RGB\((\d+),\s*(\d+),\s*(\d+)\)", color_value, re.IGNORECASE)
                        if rgb_match:
                            r, g, b = map(int, rgb_match.groups())
                            styles[current_style]["color"] = RGBColor(r, g, b)
                    else:
                        # Use predefined color map
                        color_name = color_value.lower()
                        color_map = {
                            "black": RGBColor(0, 0, 0),
                            "blue": RGBColor(0, 0, 255),
                            "red": RGBColor(255, 0, 0),
                            "green": RGBColor(0, 255, 0),
                            "yellow": RGBColor(255, 255, 0),
                            "orange": RGBColor(255, 192, 0),
                            "purple":RGBColor(182, 100, 162),
                            "gray":RGBColor(168, 173, 180),
                            
                        }
                        styles[current_style]["color"] = color_map.get(color_name, RGBColor(0, 0, 0))
                if spacing_match:
                    styles[current_style]["spacing_after"] = int(spacing_match.group(1))
                if underline_match:
                    styles[current_style]["underline"] = underline_match.group(1).strip().lower() == "yes"

        return styles

    except Exception as e:
        print(f"Error processing styling guide: {e}")
        return {}

def apply_styles(doc, text, styles):
    try:
        # Ensure styles is a dictionary even if not provided.
        if styles is None:
            styles = {}

        # Merge provided styles with default_styles.
        for style_category, defaults in default_styles.items():
            print("line 425")
            if style_category not in styles or not styles[style_category]:
                styles[style_category] = defaults
            else:
                for attr, default_value in defaults.items():
                    if attr not in styles[style_category]:
                        styles[style_category][attr] = default_value

        lines = text.split("\n")

        # Process each line.
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Title Formatting
            if line.lower().startswith("title:") or line.lower().startswith("*title:") or line.lower().startswith("%%"):
                content = line.replace("Title:", "").replace("*Title:", "").replace("%%", "").strip()
                para = doc.add_paragraph(content)
                para.alignment = styles["Title"]["alignment"]
                run = para.runs[0]
                run.bold = styles["Title"]["bold"]
                run.italic = styles["Title"]["italic"]
                run.font.size = Pt(styles["Title"]["size"])
                run.font.name = styles["Title"]["font_name"]
                run.font.color.rgb = styles["Title"]["color"]
                para.paragraph_format.space_after = Pt(styles["Title"]["spacing_after"])
                para.paragraph_format.space_before = Pt(styles["Title"]["spacing_before"])
                para.paragraph_format.line_spacing = styles["Title"]["line_spacing"]

            # Headings - Heading 1
            elif line.startswith("* ") or line.startswith("*"):
                content = line.replace("* ", "").replace("*","").strip()
                para = doc.add_paragraph(content)
                para.alignment = styles["Heading 1"]["alignment"]
                run = para.runs[0]
                run.bold = styles["Heading 1"]["bold"]
                run.italic = styles["Heading 1"]["italic"]
                run.font.size = Pt(styles["Heading 1"]["size"])
                run.font.name = styles["Heading 1"]["font_name"]
                run.font.color.rgb = styles["Heading 1"]["color"]
                para.paragraph_format.space_after = Pt(styles["Heading 1"]["spacing_after"])
                para.paragraph_format.space_before = Pt(styles["Heading 1"]["spacing_before"])
                para.paragraph_format.line_spacing = styles["Heading 1"]["line_spacing"]

            # Headings - Heading 2
            elif line.startswith("** ") or line.startswith("**"):
                content = line.replace("** ", "").replace("**","").strip()
                para = doc.add_paragraph(content)
                para.alignment = styles["Heading 2"]["alignment"]
                run = para.runs[0]
                run.bold = styles["Heading 2"]["bold"]
                run.italic = styles["Heading 2"]["italic"]
                run.font.size = Pt(styles["Heading 2"]["size"])
                run.font.name = styles["Heading 2"]["font_name"]
                run.font.color.rgb = styles["Heading 2"]["color"]
                para.paragraph_format.space_after = Pt(styles["Heading 2"]["spacing_after"])
                para.paragraph_format.space_before = Pt(styles["Heading 2"]["spacing_before"])
                para.paragraph_format.line_spacing = styles["Heading 2"]["line_spacing"]

            # Bullet Points
            elif line.startswith("• ") or line.startswith("- "):
                content = line.replace("• ", "").replace("- ", "").strip()
                para = doc.add_paragraph(content, style="List Bullet")
                run = para.runs[0]
                run.bold = styles["Bullet"]["bold"]
                run.italic = styles["Bullet"]["italic"]
                run.font.size = Pt(styles["Bullet"]["size"])
                run.font.name = styles["Bullet"]["font_name"]
                run.font.color.rgb = styles["Bullet"]["color"]
                para.paragraph_format.space_after = Pt(styles["Bullet"]["spacing_after"])
                para.paragraph_format.space_before = Pt(styles["Bullet"]["spacing_before"])
                para.paragraph_format.line_spacing = styles["Bullet"]["line_spacing"]

            # Hyperlinks
            elif re.search(r"https?://[^\s]+", line):
                para = doc.add_paragraph()
                words = line.split()
                for word in words:
                    if re.match(r"https?://[^\s]+", word):
                        run = para.add_run(word)
                        run.font.color.rgb = styles["Hyperlink"]["color"]
                        run.underline = styles["Hyperlink"]["underline"]
                        run.bold = styles["Hyperlink"]["bold"]
                        run.italic = styles["Hyperlink"]["italic"]
                        run.font.size = Pt(styles["Hyperlink"]["size"])
                        run.font.name = styles["Hyperlink"]["font_name"]
                    else:
                        para.add_run(f"{word} ")

            # Normal Paragraphs
            else:
                para = doc.add_paragraph(line)
                para.alignment = styles["Paragraph"]["alignment"]
                para.paragraph_format.space_after = Pt(styles["Paragraph"]["spacing_after"])
                para.paragraph_format.space_before = Pt(styles["Paragraph"]["spacing_before"])
                para.paragraph_format.line_spacing = styles["Paragraph"]["line_spacing"]
                run = para.runs[0]
                run.bold = styles["Paragraph"]["bold"]
                run.italic = styles["Paragraph"]["italic"]
                run.font.size = Pt(styles["Paragraph"]["size"])
                run.font.name = styles["Paragraph"]["font_name"]
                run.font.color.rgb = styles["Paragraph"]["color"]

    except Exception as e:
        print(f"Error applying styles: {e}")

# Function to format the document based on extracted styles
def format_document(styling_guide, agent_output_text, output_filename):
    try: 
        styles = read_styling_guide(styling_guide)
        print("styles: ",styles)
        doc = Document()
        apply_styles(doc, agent_output_text, styles)
        doc.save(output_filename)
        
    except Exception as e:
        print(f"Error formatting document: {e}")
  
# Function to generate edited content conflict report        
def generate_edited_content_report(edited_content, generated_content): 
    try:
        # Executing Edited Content Validation task
        validator_agent = edited_content_validator_agent()
        validator_task = edited_content_validator_task(edited_content, generated_content)

        edited_content_report = crew_result([validator_agent], [validator_task])
        
        return edited_content_report
    except Exception as e:
        print(f"Error during generate_edited_content_report: {e}")
        return "No report generated."
 
#Method to regenerate content   
def regenerate_content(edited_content, user_guidelines,user_instructions, configurations:dict, output_file_path):
    try:
        # detect guidelines and instructions language
        detect_guideline_lang = detect_language(user_guidelines)
        detect_instructions_lang = detect_language(user_instructions)
        
        #translate guidelines and instructions to english if required
        if detect_guideline_lang != "en":
            translated_guidelines = translate_text(guidelines, to_lang="en", from_lang=detect_guideline_lang)
            guidelines = translated_guidelines
            print("*******translated guidelines: ",guidelines)
        
        if detect_instructions_lang != "en":
            translated_instructions = translate_text(instructions, to_lang="en", from_lang=detect_instructions_lang)
            instructions = translated_instructions
            print("*******translated instructions: ",instructions)
        #detect edited content language
        detect_edited_content_lang = detect_language(edited_content)
        if detect_edited_content_lang != "en":
            translated_post = translate_text(edited_content, to_lang="en", from_lang=detect_edited_content_lang)
            edited_content = translated_post
        configurations_instructions = additional_instructions(configurations)
        
        combined_instructions = combined_inputs(user_guidelines, user_instructions, configurations_instructions)
        # Executing Edited Content Validation task
        regenerator_agent = content_regenerator_agent()
        regenrator_task = regenerate_content_task(edited_content, combined_instructions)
        regenerated_content = crew_result([regenerator_agent], [regenrator_task])
        
        if detect_edited_content_lang != "en":
            translated_content = translate_text(regenerated_content, to_lang=detect_edited_content_lang, from_lang="en")
            regenerated_content = translated_content
        
        regenerate_content_document(output_file_path, user_guidelines, regenerated_content)

        return regenerated_content
    except Exception as e:
        print(f"Error during write_edited_content: {e}")
        return "No content generated."
    
def image_integration(output_file_path, generate_image_flag, image_name, content, guidelines, user_instructions):
    try:
        image = None
        if generate_image_flag:
            print("generate image flag is true")
            combined_instructions = combined_inputs(guidelines, user_instructions, additional_instructions="")
            #generating prompt for image generation
            prompt_agent = create_prompt_agent()
            prompt_task = create_prompt_task(content, combined_instructions)
            prompt = crew_result([prompt_agent], [prompt_task])
            
            #Executing image generator task
            image = create_image(prompt,image_name)
            if image is not None:
                insert_image_in_doc(output_file_path, image)
        else: 
            if image_name:
                if not os.path.exists("images"):
                    os.makedirs("images") 
                image_path = os.path.join("images", image_name)
                print("image path is : ",image_path)
                insert_image_in_doc(output_file_path, image_path)
        return image
    except Exception as e:
        print(f"Error during image_integration: {e}")
        
def insert_image_in_doc(doc_path, image_path, image_width_inch=5):
    # Load the original document
    doc = Document(doc_path)
 
    # Make a backup copy of the original body elements    
    original_elements = [deepcopy(el) for el in doc._element.body]    
    
    # Remove all elements from the document body    
    for child in list(doc._element.body):    
        doc._element.body.remove(child)    
    
    # Create a new paragraph and center it  
    paragraph = doc.add_paragraph()  
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  

    # Add the image into a new run in this centered paragraph  
    run = paragraph.add_run()  
    run.add_picture(image_path, width=Inches(image_width_inch))  
    
    # Optional: Add a blank paragraph for additional spacing if desired  
    doc.add_paragraph('')  
        
    # Re-insert the original content below    
    for el in original_elements:    
        doc._element.body.append(el)    
        
    # Save the modified document    
    doc.save(doc_path) 
    
# Detect the language of input text
def detect_language(text: str) -> str:
    try:
        url = f"{AZURE_TRANSLATOR_ENDPOINT}/detect"
        params = {'api-version': API_VERSION}
        headers = {
            'Ocp-Apim-Subscription-Key': AZURE_TRANSLATOR_KEY,
            'Ocp-Apim-Subscription-Region': AZURE_REGION,
            'Content-Type': 'application/json'
        }
        body = [{'text': text}]
        response = requests.post(url, params=params, headers=headers, json=body, timeout=10, verify=False)
        response.raise_for_status()
        return response.json()[0]['language']
    except Exception as e:
        print(f"Error during detect_language: {e}")
        return "No language detected"
 
# Translate text
def translate_text(text: str, to_lang: str, from_lang: str = None) -> str:
    try:
        url = f"{AZURE_TRANSLATOR_ENDPOINT}/translate"
        params = {
            'api-version': API_VERSION,
            'to': to_lang
        }
        if from_lang:
            params['from'] = from_lang
    
        headers = {
            'Ocp-Apim-Subscription-Key': AZURE_TRANSLATOR_KEY,
            'Ocp-Apim-Subscription-Region': AZURE_REGION,
            'Content-Type': 'application/json'
        }
        body = [{'text': text}]
        response = requests.post(url, params=params, headers=headers, json=body, verify=False)
        response.raise_for_status()
        return response.json()[0]['translations'][0]['text']  
    except Exception as e:
        print(f"Error during translate_text: {e}")
        return "Translation failed" 
    
def get_lang_code(language_name: str) -> str:
    code = LANGUAGE_NAME_TO_CODE.get(language_name.strip().lower())
    if not code:
        raise ValueError(f"Unsupported or unknown language: {language_name}")
    return code

def extract_and_resize_images(docx_path, output_folder, size=(300, 300)):
    read_images_folder = os.path.join(output_folder, "read_images")
    os.makedirs(read_images_folder, exist_ok=True)
    doc = docx.Document(docx_path)
    image_count = 0

    for rel in doc.part._rels:
        rel = doc.part._rels[rel]
        if "image" in rel.target_ref:
            image_count += 1
            img_data = rel.target_part.blob
            img_ext = rel.target_part.content_type.split('/')[-1]

            # Open the image with PIL
            img = Image.open(BytesIO(img_data))

            # Resize the image
            resized_img = img.resize(size)

            # Save the resized image in the read_images folder
            img_name = f"image_{image_count}.{img_ext}"
            img_path = os.path.join(read_images_folder, img_name)
            resized_img.save(img_path)

    return read_images_folder


def delete_read_images_folder(folder_path):
    """
    Deletes the specified folder and all its contents.

    :param folder_path: Path to the folder to delete.
    """
    if os.path.exists(folder_path):
        shutil.rmtree(folder_path)
        print(f"Deleted folder: {folder_path}")
    else:
        print(f"Folder '{folder_path}' does not exist.")
        
def regenerate_content_document(docx_path, styling_guide, content):
    """
    Finalizes the document by applying the styling guide and formatting the text.  
    """
    
    print("reading the images")
    
    #read images from document
    read_images_folder_path = extract_and_resize_images(docx_path, "images", size=(1024, 1024))
    print("images read and saved at : ",read_images_folder_path)
    
    print("inserting edited content in document")
    format_document(styling_guide,content, docx_path)
    print("inserted edited content to location: ",docx_path)
    
    #iterate through the read images folder path to access all images
    for image_name in os.listdir(read_images_folder_path):
        print("inserting image in document: ",image_name)
        image_path = os.path.join(read_images_folder_path, image_name)
        # Add the image to the document
        insert_image_in_doc(docx_path, image_path)
        
    # delete the images folder
    print("deleting the images folder")
    delete_read_images_folder(read_images_folder_path)
    print("images folder deleted")
    
    print(f"Document saved with applied styles and images: {docx_path}")